"""Minimal .docx parser — paragraphs + headings + tables."""

from pathlib import Path

from harness.ingestion.models import ParsedDoc


def parse_docx(path: Path, tenant_id: str | None = None) -> ParsedDoc:
    try:
        from docx import Document as DocxDocument
    except ImportError:
        raise ImportError(
            "python-docx is required for DOCX parsing. "
            "Install: pip install python-docx"
        )
    doc = DocxDocument(str(path))
    lines: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            style = para.style.name.lower() if para.style else ""
            if "heading" in style:
                lines.append(f"# {text}")
            else:
                lines.append(text)
    # Tables as text blocks
    for i, table in enumerate(doc.tables):
        lines.append(f"\n[Table {i + 1}]")
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            lines.append(" | ".join(cells))
    return ParsedDoc(
        source_filename=path.name,
        content_type="docx",
        text="\n".join(lines),
        metadata={"source": "python-docx"},
        tenant_id=tenant_id,
    )
